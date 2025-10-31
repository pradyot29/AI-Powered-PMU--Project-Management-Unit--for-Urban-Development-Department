from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import os
from datetime import datetime
from app.core.config import settings

class DocumentGenerator:
    def __init__(self):
        self.output_dir = settings.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_docx(self, content: str, doc_type: str, doc_id: str) -> str:
        """Generate DOCX file from content"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)
        
        # Add header if needed
        header_text = f"{settings.DEPARTMENT_NAME}\n{settings.STATE_NAME}"
        header = doc.add_paragraph(header_text)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_format = header.runs[0].font
        header_format.bold = True
        header_format.size = Pt(14)
        
        doc.add_paragraph()  # Spacing
        
        # Add main content
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para_format = para.paragraph_format
                para_format.line_spacing = 1.5
                para_format.space_after = Pt(6)
                
                # Make certain lines bold (headers, subjects, etc.)
                if any(keyword in para_text.upper() for keyword in ['SUBJECT:', 'FROM:', 'TO:', 'DATE:', 'REFERENCE', 'MINUTES OF MEETING']):
                    para.runs[0].font.bold = True
        
        # Save file
        filename = f"{doc_type}_{doc_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_pdf(self, content: str, doc_type: str, doc_id: str) -> str:
        """Generate PDF file from content"""
        filename = f"{doc_type}_{doc_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4,
                              rightMargin=72, leftMargin=90,
                              topMargin=72, bottomMargin=18)
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading1'],
            fontSize=14,
            textColor='black',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            textColor='black',
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            fontName='Helvetica',
            leading=16
        )
        
        bold_style = ParagraphStyle(
            'CustomBold',
            parent=body_style,
            fontName='Helvetica-Bold'
        )
        
        # Build document
        story = []
        
        # Add header
        header_text = f"<b>{settings.DEPARTMENT_NAME}</b><br/>{settings.STATE_NAME}"
        story.append(Paragraph(header_text, header_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add content
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check if paragraph should be bold
                if any(keyword in para_text.upper() for keyword in ['SUBJECT:', 'FROM:', 'TO:', 'DATE:', 'REFERENCE', 'MINUTES OF MEETING']):
                    para_text = f"<b>{para_text}</b>"
                
                story.append(Paragraph(para_text.strip().replace('\n', '<br/>'), body_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        return filepath
    
    def generate_both_formats(self, content: str, doc_type: str, doc_id: str) -> dict:
        """Generate both PDF and DOCX"""
        try:
            docx_path = self.generate_docx(content, doc_type, doc_id)
            pdf_path = self.generate_pdf(content, doc_type, doc_id)
            
            return {
                "docx_path": docx_path,
                "pdf_path": pdf_path,
                "success": True
            }
        except Exception as e:
            return {
                "docx_path": None,
                "pdf_path": None,
                "success": False,
                "error": str(e)
            }